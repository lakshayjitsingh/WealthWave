from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.style import WD_STYLE_TYPE
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

OUTPUT = r"C:\Users\Owner\Downloads\WealthWave_Report.docx"

doc = Document()

# --- Page margins ---
for section in doc.sections:
    section.top_margin    = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin   = Inches(1.25)
    section.right_margin  = Inches(1.25)

def h(text, level=1):
    p = doc.add_heading(text, level=level)
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    return p

def para(text, bold=False, italic=False, size=12, center=False):
    p = doc.add_paragraph()
    if center:
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(text)
    run.bold = bold
    run.italic = italic
    run.font.size = Pt(size)
    return p

def bullet(items):
    for item in items:
        p = doc.add_paragraph(item, style='List Bullet')
        p.runs[0].font.size = Pt(12)

def page_break():
    doc.add_page_break()

# ===================== TITLE PAGE =====================
para("An end semester project report on", center=True, size=13)
para("", size=6)
para("WEALTHWAVE AI", bold=True, center=True, size=22)
para("A Cloud-Native Personal Finance Platform with Real-Time AI Insights", italic=True, center=True, size=13)
para("", size=6)
para("as a part of practical for", center=True, size=12)
para("", size=6)
para("Program: B.Tech (AI & ML)\nCourse: Capstone Project (CSE499)\nDepartment: AIT-CSE\nSection & Group: [Your Section & Group]", center=True, size=12)
para("", size=6)
para("submitted by", center=True, size=12)
para("Lakshayjit Singh (24BAI70225)", bold=True, center=True, size=13)
para("under guidance of", center=True, size=12)
para("[Faculty Name] ([EID])\n[Faculty Designation]", italic=True, center=True, size=12)
para("", size=6)
para("Chandigarh University, Mohali | Academic Session 2024 - 2025", center=True, size=12)
page_break()

# ===================== CERTIFICATE =====================
h("Certificate", level=1)
para('This is to certify that the project report titled "WealthWave AI — A Cloud-Native Personal Finance Platform", submitted by Lakshayjit Singh, Roll No. 24BAI70225, in partial fulfilment of the requirements of the End-Semester Evaluation, is a record of original work carried out under my supervision.')
para("")
para("Faculty Guide")
para("Name: [Faculty's Name]")
para("Signature: __________")
page_break()

# ===================== DECLARATION =====================
h("Declaration", level=1)
para('I, Lakshayjit Singh (24BAI70225), hereby declare that the project titled "WealthWave AI" is an original piece of work carried out by me as part of the requirements of the course CSE499 - Capstone Project, under the Department of AIT-CSE, Chandigarh University. I affirm that this report has not been submitted in any form for any other course or examination, and that all information, diagrams, and references used have been duly acknowledged as per academic guidelines.')
para("")
para("Lakshayjit Singh")
para("Signature: __________")
page_break()

# ===================== ACKNOWLEDGEMENT =====================
h("Acknowledgement", level=1)
para("I would like to express my sincere gratitude to my project guide for their constant support, guidance, and valuable suggestions throughout the development of this project. Their encouragement helped me stay focused and improve the quality of my work.")
para("I also thank the Department of AIT-CSE, Chandigarh University, for providing the necessary resources and a conducive learning environment that enabled me to apply classroom knowledge to a practical production-grade system.")
para("Lastly, I am grateful to my friends and family for their continuous support, motivation, and understanding during the completion of this project.")
page_break()

# ===================== ABSTRACT =====================
h("Abstract", level=1)
para("Managing personal finances efficiently is a growing challenge for individuals across age groups. Most existing tools are either too complex, require paid subscriptions, or fail to synchronize data across devices. WealthWave AI addresses these problems by providing a cloud-native, cross-platform personal finance management system.")
para("The system allows users to track income and expenses, set savings goals, estimate taxes for multiple countries, and receive real-time AI-powered financial advice through an integrated Google Gemini-2.0 chatbot. The platform is built using React 19 (frontend), Node.js/Express (backend), and PostgreSQL hosted on Neon Cloud, with deployment on Vercel (web) and Render (backend). It is also available as a native Windows Desktop application built with Electron.js.")
para("Results demonstrate successful cloud migration, live global deployment, and seamless cross-platform data synchronization, making WealthWave AI a complete, production-ready financial solution.")
page_break()

# ===================== TABLE OF CONTENTS =====================
h("Table of Contents", level=1)
toc_items = [
    "Certificate", "Declaration", "Acknowledgement", "Abstract",
    "Chapter 1: Introduction", "  1.1 Background", "  1.2 Purpose",
    "  1.3 Objectives", "  1.4 Scope", "  1.5 Target Users", "  1.6 Technology Stack Overview",
    "Chapter 2: Methodology", "  2.1 Literature Survey", "  2.2 Comparison with Existing Systems",
    "  2.3 Drawbacks of Current Systems",
    "Chapter 3: Problem Definition & Proposed Solution", "  3.1 Identified Problems",
    "  3.2 Core Requirements", "  3.3 Proposed System Overview",
    "Chapter 4: Requirement Analysis", "  4.1 Functional Requirements",
    "  4.2 Non-Functional Requirements", "  4.3 System Requirements",
    "Chapter 5: System Design", "  5.1 System Architecture", "  5.2 Database Schema",
    "  5.3 Module Descriptions",
    "Chapter 6: Implementation", "  6.1 Technology Stack", "  6.2 Module Descriptions",
    "  6.3 UI Screenshots",
    "Chapter 7: Testing", "  7.1 Test Plan", "  7.2 Test Cases", "  7.3 Error Handling",
    "Chapter 8: Challenges Faced & Solutions",
    "Chapter 9: Results & Discussion",
    "Chapter 10: Conclusion & Future Scope",
    "References",
]
for item in toc_items:
    p = doc.add_paragraph(item, style='List Paragraph')
    p.runs[0].font.size = Pt(12)
page_break()

# ===================== CHAPTER 1: INTRODUCTION =====================
h("Chapter 1", level=1)
h("Introduction", level=2)

h("1.1 Background", level=3)
para("In today's fast-paced world, individuals increasingly struggle to manage their finances effectively. Traditional methods such as spreadsheets and notebooks are error-prone, difficult to maintain, and inaccessible across multiple devices. While several digital finance tools exist, they are often overly complex, paywalled, or fail to provide intelligent insights tailored to individual users.")
para("WealthWave AI was created to solve this real-world problem by combining the power of cloud computing, artificial intelligence, and modern web development into a single unified platform. The application is designed for college students, young professionals, and everyday users who want a simple yet powerful tool to take control of their money.")

h("1.2 Purpose", level=3)
para("The purpose of WealthWave AI is to provide a free, accessible, and intelligent personal finance management system that works seamlessly across web browsers and Windows desktop applications. The platform enables users to track every rupee they earn or spend, understand their financial health through visual analytics, and receive personalized advice from an AI financial advisor.")

h("1.3 Objectives", level=3)
bullet([
    "Build a cloud-hosted, cross-platform financial management system accessible from any device",
    "Integrate a real-time AI advisor (Google Gemini-2.0) that gives personalized financial insights",
    "Implement a secure authentication system using JWT tokens and Bcrypt password hashing",
    "Provide a multi-currency dashboard with live exchange rates and tax estimation for 4 countries",
    "Package the application as a native Windows Desktop App using Electron.js",
    "Enable PDF and CSV export of financial reports for record-keeping",
])

h("1.4 Scope", level=3)
para("WealthWave AI covers the following scope:")
bullet([
    "User registration, authentication, and session management",
    "Income and expense transaction tracking with categories",
    "Real-time dashboard with charts (Cash Flow and Expense Breakdown)",
    "AI-powered financial advice chatbot using Google Gemini-2.0",
    "Savings goals creation and progress tracking",
    "Tax estimation for India, USA, UK, and Germany",
    "Multi-currency support with live exchange rates",
    "PDF and CSV report generation and download",
    "Native Windows desktop application distribution",
])

h("1.5 Target Users", level=3)
bullet([
    "College students managing limited monthly budgets",
    "Young professionals tracking salary and expenses",
    "Freelancers monitoring variable income and project payments",
    "Any individual who wants a simple, intelligent finance dashboard",
])

h("1.6 Technology Stack Overview", level=3)
table = doc.add_table(rows=8, cols=2)
table.style = 'Table Grid'
headers = ["Layer", "Technology"]
data = [
    ("Frontend", "React 19, Vite, React Router, Recharts"),
    ("Styling", "Vanilla CSS, Glassmorphism Design"),
    ("Desktop App", "Electron.js (Windows .exe Installer)"),
    ("Backend", "Node.js, Express.js, JWT, Bcrypt"),
    ("Database", "PostgreSQL (Neon Cloud)"),
    ("Deployment", "Vercel (Frontend) + Render (Backend)"),
    ("AI Engine", "Google Gemini-2.0 via OpenRouter API"),
]
for i, (k, v) in enumerate([("Layer","Technology")] + data):
    row = table.rows[i]
    row.cells[0].text = k
    row.cells[1].text = v
    for cell in row.cells:
        cell.paragraphs[0].runs[0].font.size = Pt(11)
        if i == 0:
            cell.paragraphs[0].runs[0].bold = True
page_break()

# ===================== CHAPTER 2: METHODOLOGY =====================
h("Chapter 2", level=1)
h("Methodology", level=2)

h("2.1 Literature Survey", level=3)
para("Several existing personal finance tools were studied prior to the development of WealthWave AI:")
bullet([
    "Mint (Intuit): A popular US-based finance tracker. While feature-rich, it is not available in India and requires bank account linking, which raises privacy concerns.",
    "Walnut / ET Money: Indian finance apps with UPI integration. However, they lack AI-driven advice and cross-platform desktop support.",
    "Microsoft Excel: Widely used for manual budgeting. Lacks real-time updates, AI assistance, and cloud synchronization.",
    "Google Sheets: Collaborative but requires manual data entry and offers no financial intelligence or automated analysis.",
])
para("The literature review concluded that no existing free tool combines cloud sync, AI advice, multi-currency support, and desktop app distribution in a single platform — the gap that WealthWave AI fills.")

h("2.2 Comparison with Existing Systems", level=3)
table2 = doc.add_table(rows=5, cols=5)
table2.style = 'Table Grid'
comp_data = [
    ("Feature", "WealthWave AI", "Mint", "ET Money", "Excel"),
    ("AI Financial Advice", "Yes", "No", "No", "No"),
    ("Cloud Sync", "Yes (Neon DB)", "Yes", "Yes", "No"),
    ("Desktop App", "Yes (Electron)", "No", "No", "Yes"),
    ("Multi-Currency + Tax", "Yes", "No", "No", "Manual"),
]
for i, row_data in enumerate(comp_data):
    row = table2.rows[i]
    for j, cell_text in enumerate(row_data):
        row.cells[j].text = cell_text
        run = row.cells[j].paragraphs[0].runs[0] if row.cells[j].paragraphs[0].runs else row.cells[j].paragraphs[0].add_run(cell_text)
        run.font.size = Pt(10)
        if i == 0:
            run.bold = True

h("2.3 Drawbacks of Current Systems", level=3)
bullet([
    "Most tools are region-locked (USA only) or have limited international support",
    "No existing free tool provides an integrated AI financial advisor with context awareness",
    "Existing apps require invasive bank account access to function",
    "Desktop and web versions of existing tools often show different data (no true sync)",
])
page_break()

# ===================== CHAPTER 3: PROBLEM DEFINITION =====================
h("Chapter 3", level=1)
h("Problem Definition & Proposed Solution", level=2)

h("3.1 Identified Problems", level=3)
bullet([
    "Data Isolation: Financial data stored locally (SQLite/files) is lost or inaccessible when switching devices",
    "No Intelligence: Users have no AI-driven system to analyze spending patterns or suggest savings strategies",
    "Platform Gap: Desktop apps and web apps don't stay in sync — users see different data on each platform",
    "Manual Tax Calculation: No built-in tool for estimating taxes across different countries for Indian users",
    "Poor UX: Most finance apps have outdated, complex interfaces that discourage regular use",
])

h("3.2 Core Requirements", level=3)
bullet([
    "A single cloud database that serves both the web and desktop applications",
    "An AI chatbot that understands the user's actual financial data, not generic advice",
    "Secure login system that protects user data with industry-standard encryption",
    "A beautiful, modern UI that makes daily financial tracking enjoyable",
    "A downloadable Windows app for users who prefer a desktop experience",
])

h("3.3 Proposed System Overview", level=3)
para("WealthWave AI proposes a 3-tier cloud architecture:")
bullet([
    "Tier 1 (Presentation): React 19 web app on Vercel + Electron.js Windows desktop app",
    "Tier 2 (Logic): Node.js/Express REST API on Render with JWT authentication and AI bridge",
    "Tier 3 (Data): PostgreSQL database on Neon Cloud with SSL-encrypted connections",
])
para("This architecture ensures that any change made on the website is instantly visible on the desktop app, and vice versa, since both talk to the same cloud database.")
page_break()

# ===================== CHAPTER 4: REQUIREMENTS =====================
h("Chapter 4", level=1)
h("Requirement Analysis", level=2)

h("4.1 Functional Requirements", level=3)
bullet([
    "FR1: Users can register a new account with name, email, and password",
    "FR2: Users can log in and receive a JWT token valid for 7 days",
    "FR3: Users can add income and expense transactions with amount, name, and category",
    "FR4: Dashboard displays real-time Cash Flow bar chart and Expense Breakdown pie chart",
    "FR5: Users can set and track Savings Goals with a visual progress meter",
    "FR6: AI chatbot responds to financial questions using the user's actual data as context",
    "FR7: Tax estimator calculates estimated annual tax for India, USA, UK, and Germany",
    "FR8: Users can switch currency (INR, USD, GBP, EUR) with live conversion rates",
    "FR9: Users can download their transaction history as PDF or CSV",
    "FR10: The system is accessible on web browsers and as a Windows desktop application",
])

h("4.2 Non-Functional Requirements", level=3)
bullet([
    "Performance: API responses must return within 500ms under normal load",
    "Security: All passwords are hashed with Bcrypt (cost factor 10); tokens use HS256 JWT",
    "Availability: Hosted on cloud services (Vercel + Render) with 99.9% uptime SLA",
    "Scalability: PostgreSQL connection pooling handles multiple concurrent users",
    "Usability: UI designed with Glassmorphism principles for a premium, intuitive experience",
])

h("4.3 System Requirements", level=3)
para("For Web Users:")
bullet(["Any modern browser (Chrome, Firefox, Edge, Safari)", "Internet connection"])
para("For Desktop App Users:")
bullet(["Windows 10 or above (64-bit)", "Minimum 4GB RAM, 700MB disk space", "Internet connection for cloud sync and AI features"])
page_break()

# ===================== CHAPTER 5: SYSTEM DESIGN =====================
h("Chapter 5", level=1)
h("System Design", level=2)

h("5.1 System Architecture", level=3)
para("WealthWave AI follows a 3-Tier Cloud Architecture:")
para("[ User Device ] → HTTPS → [ Vercel: React Frontend ] → REST API → [ Render: Node.js Backend ] → SQL → [ Neon: PostgreSQL DB ]", italic=True)
para("The Electron.js desktop app bundles the React frontend and communicates with the same Render backend, ensuring full data parity between web and desktop users.")

h("5.2 Database Schema", level=3)
para("The PostgreSQL database contains the following tables:")
bullet([
    "users: id, name, email, password (hashed), budget, created_at",
    "transactions: id, user_id (FK), type, amount, name, category, date",
    "goals: id, user_id (FK), name, target_amount, saved_amount, created_at",
])

h("5.3 Module Descriptions", level=3)
bullet([
    "Auth Module: Handles /register and /login. Validates input, hashes passwords with Bcrypt, issues JWT tokens",
    "Transaction Module: GET /transactions returns all user transactions; POST /transactions adds a new one",
    "Dashboard Module: GET /summary returns aggregated income, expense, and balance totals",
    "Goals Module: Full CRUD (Create, Read, Update, Delete) for user savings goals",
    "AI Module: POST /ai/insights — sends user's financial context to Gemini-2.0 and returns advice",
])
page_break()

# ===================== CHAPTER 6: IMPLEMENTATION =====================
h("Chapter 6", level=1)
h("Implementation", level=2)

h("6.1 Technology Stack", level=3)
bullet([
    "Frontend: React 19 with Vite build tool. React Router for navigation. Recharts for animated charts.",
    "Backend: Node.js with Express.js. Async/await pattern for all database queries. CORS configured for Vercel domain.",
    "Database: PostgreSQL on Neon.tech with connection pooling via the 'pg' library.",
    "Authentication: JWT tokens (7-day expiry) + Bcrypt password hashing (cost factor 10).",
    "AI Integration: Axios POST to OpenRouter API with Google Gemini-2.0 model. HTTP-Referer and X-Title headers included.",
    "Desktop: Electron.js wraps the production Vite build into a native Windows application.",
    "Deployment: Git push to GitHub triggers auto-deploy on Vercel (frontend) and Render (backend).",
])

h("6.2 Key Code Implementation", level=3)
para("Backend — PostgreSQL Query Example (async/await):", bold=True)
p = doc.add_paragraph()
p.add_run("const result = await pool.query(\n  'SELECT * FROM transactions WHERE user_id = $1 ORDER BY date DESC',\n  [req.user.id]\n);\nres.json(result.rows);").font.name = 'Courier New'

para("AI Route — OpenRouter Integration:", bold=True)
p2 = doc.add_paragraph()
p2.add_run("const response = await axios.post('https://openrouter.ai/api/v1/chat/completions', {\n  model: 'google/gemini-2.0-flash-001',\n  messages: [{ role: 'system', content: systemPrompt }, { role: 'user', content: message }]\n}, { headers: { Authorization: `Bearer ${process.env.OPENROUTER_API_KEY}`, 'HTTP-Referer': 'https://wealthwave-gamma.vercel.app' } });").font.name = 'Courier New'

h("6.3 UI Screenshots", level=3)
para("[ Insert Screenshot 1: Login Page — wealth-wave-gamma.vercel.app ]", italic=True)
para("[ Insert Screenshot 2: Main Dashboard with Charts ]", italic=True)
para("[ Insert Screenshot 3: AI Chat Panel ]", italic=True)
para("[ Insert Screenshot 4: Savings Goals Section ]", italic=True)
para("[ Insert Screenshot 5: Windows Desktop App ]", italic=True)
page_break()

# ===================== CHAPTER 7: TESTING =====================
h("Chapter 7", level=1)
h("Testing", level=2)

h("7.1 Test Plan", level=3)
para("All features were tested manually on the live production environment (Vercel + Render + Neon) using two different laptops to simulate real-world multi-device usage.")

h("7.2 Test Cases", level=3)
table3 = doc.add_table(rows=7, cols=3)
table3.style = 'Table Grid'
tc_data = [
    ("Test Case", "Input", "Expected Result"),
    ("TC01: User Registration", "Name, Email, Password", "Account created, redirected to dashboard"),
    ("TC02: User Login", "Valid email + password", "JWT token issued, dashboard loads"),
    ("TC03: Add Transaction", "Amount, name, category", "Transaction saved, chart updates"),
    ("TC04: AI Chat", "Type: 'How am I doing?'", "AI responds with personalized financial summary"),
    ("TC05: Download Report", "Click 'Download PDF'", "PDF file downloaded with all transactions"),
    ("TC06: Cross-Device Sync", "Login on 2nd laptop", "Same data visible on both devices"),
]
for i, row_data in enumerate(tc_data):
    row = table3.rows[i]
    for j, val in enumerate(row_data):
        row.cells[j].text = val
        if row.cells[j].paragraphs[0].runs:
            run = row.cells[j].paragraphs[0].runs[0]
        else:
            run = row.cells[j].paragraphs[0].add_run(val)
        run.font.size = Pt(10)
        if i == 0:
            run.bold = True

h("7.3 Error Handling", level=3)
bullet([
    "Invalid login credentials return a 401 Unauthorized response with a clear error message",
    "Missing JWT token returns 403 Forbidden on all protected routes",
    "Database connection errors are caught with try/catch and return 500 Internal Server Error",
    "OpenRouter API failures return a user-friendly error message in the chat panel",
])
page_break()

# ===================== CHAPTER 8: CHALLENGES =====================
h("Chapter 8", level=1)
h("Challenges Faced & Solutions", level=2)

challenges = [
    ("Database Migration (SQLite to PostgreSQL)",
     "The original app used SQLite with callback-based queries. Migrating to PostgreSQL required rewriting all routes to use async/await with parameterized queries ($1, $2). Solved by refactoring each route file systematically."),
    ("Vercel Deployment Failure (Dependency Conflict)",
     "Vercel's npm install failed due to a peer dependency conflict between Vite 8.x and vite-plugin-pwa. Solved by setting the Install Command to 'npm install --legacy-peer-deps' in Vercel's build settings."),
    ("AI Rejection in Production",
     "OpenRouter rejected API requests from the live Render server. Root cause: missing HTTP-Referer and X-Title headers required for production environments. Solved by adding these headers to the axios request."),
    ("Windows App Using Old Localhost URLs",
     "The distributed .exe was built before the cloud migration, so it pointed to localhost:5000 instead of the Render URL. Solved by rebuilding the Electron app after updating all API URLs, and re-uploading to Google Drive."),
    ("Cold Start Latency on Render Free Tier",
     "The Render free tier spins down the server after 15 minutes of inactivity, causing a ~15 second delay on first load. This is a known limitation of free hosting tiers; a paid tier would eliminate it."),
]
for title, solution in challenges:
    h(title, level=3)
    para(solution)
page_break()

# ===================== CHAPTER 9: RESULTS =====================
h("Chapter 9", level=1)
h("Results & Discussion", level=2)

h("9.1 Key Output Screens", level=3)
para("The application was successfully deployed and tested on multiple devices. Key screens include the Login/Register page, the main financial dashboard, the AI chatbot panel, and the savings goals tracker. Screenshots are included in Section 6.3.")

h("9.2 Performance", level=3)
bullet([
    "Frontend (Vercel): Average page load time < 2 seconds via global CDN",
    "Backend (Render): Average API response time < 500ms",
    "AI Response: Average Gemini-2.0 response time: 2-4 seconds",
    "Database: Neon PostgreSQL handles all queries with connection pooling",
    "Desktop App: 611 MB installer, single-click installation on Windows 10/11",
])

h("9.3 Learning Outcomes", level=3)
bullet([
    "Gained hands-on experience with full-stack cloud deployment using industry tools",
    "Learned PostgreSQL async/await patterns and database migration strategies",
    "Understood JWT authentication and Bcrypt hashing for production security",
    "Integrated a real AI API (Google Gemini-2.0) into a live production application",
    "Packaged a React app as a distributable Windows application using Electron.js",
    "Managed CI/CD through GitHub, with auto-deployment to Vercel and Render",
])
page_break()

# ===================== CHAPTER 10: CONCLUSION =====================
h("Chapter 10", level=1)
h("Conclusion & Future Scope", level=2)

h("Conclusion", level=3)
para("WealthWave AI successfully demonstrates the integration of cloud computing, artificial intelligence, and modern web development to solve a real-world personal finance problem. The project evolved from a local SQLite-based application to a fully production-ready, globally accessible platform.")
para("Key achievements include:")
bullet([
    "Successfully migrated from local SQLite to Neon PostgreSQL cloud database",
    "Deployed a live, globally accessible web application on Vercel",
    "Integrated Google Gemini-2.0 AI for context-aware financial advice",
    "Packaged and distributed a native Windows desktop application",
    "Implemented industry-standard security (JWT + Bcrypt + SSL)",
])
para("The primary limitation is the cold-start delay on the Render free tier. A paid hosting plan would eliminate this issue entirely.")

h("Future Scope", level=3)
bullet([
    "Android & iOS Mobile App using React Native with the same cloud backend",
    "UPI / Bank Integration — Auto-import transactions from bank statements",
    "Investment Portfolio Tracker — Monitor Stocks, Mutual Funds & Crypto",
    "Predictive AI — ML model to forecast future expenses and suggest saving targets",
    "Family Budgeting Mode — Shared financial dashboard for household management",
    "Smart Alerts — Push notifications for unusual spending and budget breaches",
])
page_break()

# ===================== REFERENCES =====================
h("References", level=1)
refs = [
    "React Documentation. (2024). React 19 Official Docs. https://react.dev",
    "Node.js Foundation. (2024). Node.js Official Documentation. https://nodejs.org/docs",
    "PostgreSQL Global Development Group. (2024). PostgreSQL 16 Documentation. https://www.postgresql.org/docs",
    "Neon Technologies. (2024). Neon Serverless Postgres Documentation. https://neon.tech/docs",
    "Electron Authors. (2024). Electron.js Official Documentation. https://www.electronjs.org/docs",
    "OpenRouter. (2024). OpenRouter API Reference. https://openrouter.ai/docs",
    "Google DeepMind. (2024). Gemini 2.0 Flash Model. https://deepmind.google/technologies/gemini",
    "Vercel Inc. (2024). Vercel Deployment Documentation. https://vercel.com/docs",
    "Render Inc. (2024). Render Cloud Hosting Documentation. https://docs.render.com",
    "Auth0. (2024). Introduction to JSON Web Tokens. https://jwt.io/introduction",
]
for i, ref in enumerate(refs):
    p = doc.add_paragraph(f"{i+1}. {ref}", style='List Paragraph')
    p.runs[0].font.size = Pt(12)
    p.runs[0].font.name = 'Times New Roman'

# ===================== SAVE =====================
doc.save(OUTPUT)
print("SUCCESS! Report saved to: " + OUTPUT)
